"""DOCX and plain text document parsing."""
from __future__ import annotations

from pathlib import Path

import pandas as pd

from .utils import ParsedDocument, ProcessingError, table_to_dataframe

try:
    from docx import Document
except ImportError:  # pragma: no cover - depends on local environment
    Document = None


def parse_document_file(file_path: str | Path) -> ParsedDocument:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix in {'.txt', ''}:
        try:
            text = path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            text = path.read_text(encoding='latin-1')
        return ParsedDocument(file_path=path, file_type='text', raw_text=text, tables=[], metadata={})

    if suffix != '.docx':
        raise ProcessingError(f"Unsupported document type '{suffix}'.")

    if Document is None:
        raise ProcessingError('DOCX support is unavailable. Install python-docx.')

    try:
        document = Document(str(path))
    except Exception as exc:  # pragma: no cover - library behavior
        raise ProcessingError(f"Unable to open Word document '{path.name}': {exc}") from exc

    text_parts: list[str] = []
    tables: list[pd.DataFrame] = []
    table_contexts: list[str] = []

    paragraphs = list(document.paragraphs)
    doc_tables = list(document.tables)
    paragraph_index = 0
    table_index = 0
    current_context = ''

    for child in document.element.body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            paragraph = paragraphs[paragraph_index]
            paragraph_index += 1
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
                current_context = text
        elif tag == 'tbl':
            table = doc_tables[table_index]
            table_index += 1
            raw_rows = []
            for row in table.rows:
                raw_rows.append([cell.text.strip() for cell in row.cells])
            frame = table_to_dataframe(raw_rows)
            if not frame.empty:
                frame.attrs['context_text'] = current_context
                frame.attrs['table_index'] = len(tables) + 1
                tables.append(frame)
                table_contexts.append(current_context)

    if not text_parts and not tables:
        raise ProcessingError(f"No readable text or tables found in '{path.name}'.")

    return ParsedDocument(
        file_path=path,
        file_type='docx',
        raw_text='\n'.join(text_parts),
        tables=tables,
        metadata={'table_count': len(tables), 'table_contexts': table_contexts},
    )
