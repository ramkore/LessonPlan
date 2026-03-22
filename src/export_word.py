"""Word document export for lesson plans using python-docx."""
from __future__ import annotations

from pathlib import Path

from .lesson_plan_format import (
    LAB_MAIN_COLUMN_WIDTHS,
    SUMMARY_TABLE_COLUMN_WIDTHS,
    THEORY_MAIN_COLUMN_WIDTHS,
    formatted_lesson_plan_frame,
    lab_date_row_spans,
    summary_values,
    week_row_spans,
)
from .utils import LessonPlanBundle, ProcessingError, header_image_path, report_output_paths

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError:  # pragma: no cover - depends on local environment
    Document = None


SIGN_LINE_TOP_GAP = 24


def export_bundle_to_word(bundle: LessonPlanBundle, output_dir: str | Path) -> Path:
    if Document is None:
        raise ProcessingError('Word export requires python-docx.')

    document_path = report_output_paths(bundle, output_dir)['lesson_plan_word']

    document = Document()
    _configure_document(document)
    _add_header_image(document)

    summary = summary_values(bundle)
    lab_mode = bool(summary['is_lab'])

    _add_centered_paragraph(document, summary['department_name'], size=12, bold=True)
    document.add_paragraph('')
    _add_centered_paragraph(document, summary['plan_title'], size=14, bold=True)
    document.add_paragraph('')

    summary_table = document.add_table(rows=3, cols=7)
    summary_table.style = 'Table Grid'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.autofit = False
    _set_column_widths(summary_table, SUMMARY_TABLE_COLUMN_WIDTHS)
    _populate_summary_table(summary_table, summary)
    document.add_paragraph('')

    if lab_mode:
        _append_lab_table(document, bundle)
        _add_sign_line_gap(document)
        _add_sign_line(document, summary['sign_labels'])
    else:
        _add_sign_line_gap(document)
        _add_sign_line(document, summary['sign_labels'])
        document.add_paragraph('')
        _append_regulation_row(document, summary)
        _append_theory_table(document, bundle)

    document.add_paragraph('')
    _add_books_section(document, 'TEXT BOOKS:', bundle.metadata.get('text_books', []))
    document.add_paragraph('')
    _add_books_section(document, 'REFERENCE BOOKS:', bundle.metadata.get('reference_books', []))

    document.save(document_path)
    return document_path


def _append_theory_table(document, bundle: LessonPlanBundle) -> None:
    frame = formatted_lesson_plan_frame(bundle.lesson_plan, lab_mode=False)
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_column_widths(table, THEORY_MAIN_COLUMN_WIDTHS)

    for index, header in enumerate(frame.columns):
        cell = table.rows[0].cells[index]
        cell.text = header
        _style_cell(cell, font_size=9, bold=True, center=True)

    for row in frame.to_dict('records'):
        cells = table.add_row().cells
        for index, header in enumerate(frame.columns):
            cells[index].text = str(row.get(header, ''))
            _style_cell(cells[index], font_size=8.5, center=index != 3)

    for start_index, end_index, _label in week_row_spans(frame):
        if end_index <= start_index:
            continue
        start_cell = table.cell(start_index + 1, 1)
        end_cell = table.cell(end_index + 1, 1)
        start_cell.merge(end_cell)
        _style_cell(start_cell, font_size=8.5, center=True)


def _append_lab_table(document, bundle: LessonPlanBundle) -> None:
    frame = formatted_lesson_plan_frame(bundle.lesson_plan, lab_mode=True)
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_column_widths(table, LAB_MAIN_COLUMN_WIDTHS)

    for index, header in enumerate(frame.columns):
        cell = table.rows[0].cells[index]
        cell.text = header
        _style_cell(cell, font_size=9, bold=True, center=True)

    for row in frame.to_dict('records'):
        cells = table.add_row().cells
        for index, header in enumerate(frame.columns):
            cells[index].text = str(row.get(header, ''))
            _style_cell(cells[index], font_size=8.5, center=index != 2)

    for start_index, end_index, _label in lab_date_row_spans(frame):
        if end_index <= start_index:
            continue
        start_cell = table.cell(start_index + 1, 4)
        end_cell = table.cell(end_index + 1, 4)
        start_cell.merge(end_cell)
        _style_cell(start_cell, font_size=8.5, center=True)


def _add_header_image(document) -> None:
    image_path = header_image_path()
    if image_path is None:
        return

    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    available_width = section.page_width - section.left_margin - section.right_margin
    run.add_picture(str(image_path), width=available_width)


def _configure_document(document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Inches(0.20)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(1.29)
    section.bottom_margin = Inches(1.03)

    normal_style = document.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(10)


def _populate_summary_table(table, summary: dict[str, str]) -> None:
    header = table.rows[0].cells
    header[0].text = 'Subject\nCode'
    header[1].text = 'Subject Name'
    header[2].text = 'Class/Sem'
    header[3].text = 'Faculty Name /\nDesignation'
    header[4].text = 'Number of\nStudents'
    header[5].merge(header[6]).text = summary['total_label_display']

    row_two = table.rows[1].cells
    row_two[0].text = summary['course_code']
    row_two[1].text = summary['course_title_display']
    row_two[2].text = summary['class_sem_display']
    row_two[3].text = summary['faculty_display_summary']
    row_two[4].text = summary['student_count']
    row_two[5].text = summary['lecture_label']
    row_two[6].text = summary['tutorial_label']

    row_three = table.rows[2].cells
    row_three[0].text = summary['course_code']
    row_three[1].text = summary['course_title_display']
    row_three[2].text = summary['class_sem_display']
    row_three[3].text = summary['faculty_display_summary']
    row_three[4].text = summary['student_count']
    row_three[5].text = summary['lecture_count']
    row_three[6].text = summary['tutorial_count']

    for column in range(5):
        table.cell(1, column).merge(table.cell(2, column))

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height = Inches(0.46)
    table.rows[1].height = Inches(0.50)
    table.rows[2].height = Inches(0.40)

    for row in table.rows:
        for cell in row.cells:
            _style_cell(cell, font_size=9.2, center=True)


def _append_regulation_row(document, summary: dict[str, str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = f"Regulation: {summary['regulation']}     Academic Year: {summary['academic_year']}"
    _style_cell(table.cell(0, 0), font_size=9.5, bold=True, center=False)


def _add_sign_line_gap(document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(SIGN_LINE_TOP_GAP)


def _add_sign_line(document, labels: tuple[str, str, str]) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_column_widths(table, [2.25, 1.80, 2.45])

    row = table.rows[0]
    row.height = Inches(0.38)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    cells = row.cells
    cells[0].text = labels[0]
    cells[1].text = labels[1]
    cells[2].text = labels[2]

    for cell in cells:
        _style_cell(cell, font_size=10, bold=True, center=True)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(5)


def _add_books_section(document, title: str, entries) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10.5)

    items = [str(entry).strip() for entry in entries if str(entry).strip()]
    if not items:
        placeholder = document.add_paragraph('1. -')
        _set_run_font(placeholder.runs, 'Times New Roman', 10)
        return

    for index, entry in enumerate(items, start=1):
        paragraph = document.add_paragraph(f'{index}. {entry}')
        paragraph.paragraph_format.space_after = Pt(0)
        _set_run_font(paragraph.runs, 'Times New Roman', 10)


def _set_column_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)


def _style_cell(cell, font_size: float = 9.0, bold: bool = False, center: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.bold = bold
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)


def _add_centered_paragraph(document, text: str, size: float, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)


def _set_run_font(runs, font_name: str, font_size: float) -> None:
    for run in runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
